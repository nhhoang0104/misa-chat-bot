import typing
import os
import logging
from typing import List, Optional, Dict, Any
from pathlib import Path
from langchain_core.tools import tool

# Thiết lập logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def validate_request_params(request: dict) -> Dict[str, Any]:
    """
    Kiểm tra và validate các tham số đầu vào.
    
    Args:
        request: Dictionary chứa thông tin yêu cầu
        
    Returns:
        Dict chứa kết quả validation với keys: 'valid', 'errors', 'warnings'
    """
    errors = []
    warnings = []
    
    # Kiểm tra loại bộ đề
    loai_bode = request.get('loai_bode', '').lower().strip()
    if loai_bode not in ['trắc nghiệm', 'tự luận', 'trac nghiem', 'tu luan']:
        errors.append("Loại bộ đề phải là 'trắc nghiệm' hoặc 'tự luận'")
    
    # Kiểm tra số câu hỏi
    so_cau = request.get('so_cau')
    if so_cau is None:
        errors.append("Thiếu thông tin số câu hỏi")
    else:
        try:
            so_cau = int(so_cau)
            if so_cau <= 0:
                errors.append("Số câu hỏi phải là số dương")
            elif so_cau > 50:
                warnings.append("Số câu hỏi lớn (>50) có thể mất nhiều thời gian để tạo")
        except (ValueError, TypeError):
            errors.append("Số câu hỏi phải là một số nguyên")
    
    # Kiểm tra nội dung
    chu_de = request.get('chu_de', '').strip()
    noi_dung_sach = request.get('noi_dung_sach', '').strip()
    
    if not chu_de and not noi_dung_sach:
        errors.append("Cần cung cấp ít nhất chủ đề hoặc nội dung sách")
    
    if noi_dung_sach and len(noi_dung_sach) < 10:
        warnings.append("Nội dung sách quá ngắn, có thể ảnh hưởng đến chất lượng câu hỏi")
    
    return {
        'valid': len(errors) == 0,
        'errors': errors,
        'warnings': warnings
    }

def get_missing_params(request: dict) -> List[str]:
    """
    Kiểm tra thông tin đầu vào, trả về danh sách tham số còn thiếu.
    Các tham số cần thiết: loai_bode, so_cau
    """
    missing = []
    if 'loai_bode' not in request or not request['loai_bode']:
        missing.append('loai_bode')
    if 'so_cau' not in request or not request['so_cau']:
        missing.append('so_cau')
    
    return missing

def create_question_set(request: dict) -> Dict[str, Any]:
    """
    Hàm chính nhận yêu cầu tạo bộ đề.
    Nếu thiếu thông tin, trả về câu hỏi gợi ý cho từng tham số cần hỏi lại.
    Nếu đủ, sẽ sinh bộ câu hỏi và xuất file docx.
    
    Args:
        request: Dictionary chứa thông tin yêu cầu
        
    Returns:
        Dictionary chứa kết quả xử lý
    """
    try:
        # Kiểm tra tham số thiếu
        missing = get_missing_params(request)
        if missing:
            questions = []
            for param in missing:
                if param == 'loai_bode':
                    questions.append('Bạn muốn tạo bộ đề loại nào? (trắc nghiệm/tự luận)')
                elif param == 'so_cau':
                    questions.append('Bạn muốn tạo bộ đề với bao nhiêu câu hỏi?')
                else:
                    questions.append(f'Vui lòng cung cấp thông tin cho tham số: {param}')
            return {
                'status': 'need_more_info', 
                'missing_params': missing, 
                'questions': questions
            }
        
        # Validate dữ liệu đầu vào
        validation_result = validate_request_params(request)
        if not validation_result['valid']:
            return {
                'status': 'validation_error',
                'errors': validation_result['errors'],
                'warnings': validation_result.get('warnings', [])
            }
        
        # Log warnings nếu có
        if validation_result['warnings']:
            for warning in validation_result['warnings']:
                logger.warning(warning)
        
        # Sinh câu hỏi bằng LLM Gemini
        logger.info("Bắt đầu sinh câu hỏi từ LLM...")
        questions = generate_questions_llm(request)
        
        if not questions:
            return {
                'status': 'error', 
                'message': 'Không sinh được câu hỏi từ LLM. Vui lòng thử lại sau.'
            }
        
        # Tạo file docx
        try:
            filename = f"bo_de_{request.get('chu_de', 'general').replace(' ', '_')}.docx"
            export_questions_to_docx(questions, filename)
            logger.info(f"Đã tạo file: {filename}")
            
            return {
                'status': 'success', 
                'questions': questions, 
                'docx_file': filename,
                'question_count': len(questions)
            }
        except Exception as e:
            logger.error(f"Lỗi khi tạo file docx: {str(e)}")
            return {
                'status': 'partial_success',
                'questions': questions,
                'message': f'Đã sinh câu hỏi nhưng không thể tạo file docx: {str(e)}'
            }
            
    except Exception as e:
        logger.error(f"Lỗi trong create_question_set: {str(e)}")
        return {
            'status': 'error',
            'message': f'Đã xảy ra lỗi: {str(e)}'
        }

def generate_questions_llm(request: dict) -> List[Dict[str, Any]]:
    """
    Gọi LLM Gemini qua LangChain để sinh câu hỏi, sử dụng with_structured_output.
    
    Args:
        request: Dictionary chứa thông tin yêu cầu
        
    Returns:
        List các câu hỏi đã được sinh
    """
    try:
        from langchain_google_genai import ChatGoogleGenerativeAI
        from pydantic import BaseModel, Field
        from typing import List, Optional
        from dotenv import load_dotenv
        
        load_dotenv()
        api_key = os.getenv("GEMINI_API_KEY")
        
        if not api_key:
            logger.error("Không tìm thấy GEMINI_API_KEY trong environment variables")
            return []
        
        # Định nghĩa schema cho câu hỏi
        class Question(BaseModel):
            question_type: str = Field(description="Loại câu hỏi: 'trắc nghiệm' hoặc 'tự luận'")
            question: str = Field(description="Nội dung câu hỏi")
            choices: Optional[List[str]] = Field(default=None, description="Các đáp án cho câu trắc nghiệm")
            correct_answer: Optional[str] = Field(default=None, description="Đáp án đúng")
            explanation: Optional[str] = Field(default=None, description="Giải thích đáp án")
        
        class QuestionSet(BaseModel):
            questions: List[Question] = Field(description="Danh sách câu hỏi")
        
        # Khởi tạo model
        model = ChatGoogleGenerativeAI(
            model="gemini-2.0-flash-exp",
            temperature=0.7,
            max_tokens=None,
            timeout=60,
            max_retries=3,
            google_api_key=api_key
        ).with_structured_output(QuestionSet)
        
        # Chuẩn bị thông tin
        loai_bode = request.get('loai_bode', 'trắc nghiệm').lower().strip()
        if loai_bode in ['trac nghiem']:
            loai_bode = 'trắc nghiệm'
        elif loai_bode in ['tu luan']:
            loai_bode = 'tự luận'
            
        so_cau = int(request.get('so_cau', 5))
        chu_de = request.get('chu_de', '').strip()
        noi_dung_sach = request.get('noi_dung_sach', '').strip()
        
        # Tạo prompt chi tiết
        prompt = f"""
Bạn là một chuyên gia giáo dục, hãy tạo {so_cau} câu hỏi {loai_bode} chất lượng cao.

Thông tin:
- Chủ đề: {chu_de if chu_de else 'Chung'}
- Nội dung tham khảo: {noi_dung_sach if noi_dung_sach else 'Kiến thức cơ bản'}

Yêu cầu:
"""

        if loai_bode == 'trắc nghiệm':
            prompt += """
- Mỗi câu hỏi có 4 đáp án (A, B, C, D)
- Chỉ có 1 đáp án đúng
- Các đáp án sai phải hợp lý, không quá dễ loại bỏ
- Cung cấp giải thích ngắn gọn cho đáp án đúng
- Câu hỏi phải rõ ràng, không gây nhầm lẫn
"""
        else:
            prompt += """
- Câu hỏi mở, yêu cầu tư duy và phân tích
- Không cần đáp án cụ thể
- Có thể cung cấp hướng dẫn trả lời
- Câu hỏi phải khuyến khích suy nghĩ sâu
"""

        prompt += """
Đảm bảo:
- Câu hỏi có độ khó phù hợp
- Ngôn ngữ tiếng Việt chuẩn
- Nội dung chính xác về mặt học thuật
- Đa dạng về hình thức và góc độ tiếp cận
"""

        # Gọi LLM
        logger.info("Đang gọi LLM để sinh câu hỏi...")
        response = model.invoke(prompt)
        
        # Chuyển đổi kết quả
        questions = []
        for q in response.questions:
            question_data = {
                'question_type': loai_bode,
                'question': q.question,
                'explanation': q.explanation
            }
            
            if loai_bode == 'trắc nghiệm':
                question_data['choices'] = q.choices or []
                question_data['correct_answer'] = q.correct_answer
            
            questions.append(question_data)
        
        logger.info(f"Đã sinh thành công {len(questions)} câu hỏi")
        return questions
        
    except ImportError as e:
        logger.error(f"Lỗi import thư viện: {str(e)}")
        return []
    except Exception as e:
        logger.error(f"Lỗi khi gọi LLM: {str(e)}")
        return []

def export_questions_to_docx(questions: List[Dict[str, Any]], filename: str) -> None:
    """
    Xuất danh sách câu hỏi ra file docx với format đẹp.
    
    Args:
        questions: Danh sách câu hỏi
        filename: Tên file output
    """
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Tạo thư mục output nếu chưa có
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)
        
        # Tạo document
        doc = Document()
        
        # Tiêu đề
        heading = doc.add_heading("BỘ ĐỀ KIỂM TRA", level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thông tin bộ đề
        doc.add_paragraph(f"Số câu hỏi: {len(questions)}")
        doc.add_paragraph(f"Loại đề: {questions[0].get('question_type', 'Không xác định') if questions else 'Không xác định'}")
        doc.add_paragraph("")  # Dòng trống
        
        # Thêm câu hỏi
        for idx, q in enumerate(questions, 1):
            # Số thứ tự và câu hỏi
            question_para = doc.add_paragraph()
            question_para.add_run(f"Câu {idx}: ").bold = True
            question_para.add_run(q.get('question', ''))
            
            # Nếu là trắc nghiệm, thêm các đáp án
            if q.get('question_type') == 'trắc nghiệm' and q.get('choices'):
                choices = q.get('choices', [])
                choice_labels = ['A', 'B', 'C', 'D']
                
                for i, choice in enumerate(choices[:4]):  # Chỉ lấy tối đa 4 đáp án
                    if i < len(choice_labels):
                        choice_para = doc.add_paragraph(f"   {choice_labels[i]}. {choice}")
                
                # Đáp án đúng
                if q.get('correct_answer'):
                    answer_para = doc.add_paragraph()
                    answer_para.add_run("Đáp án: ").bold = True
                    answer_para.add_run(q.get('correct_answer'))
                
                # Giải thích (nếu có)
                if q.get('explanation'):
                    explain_para = doc.add_paragraph()
                    explain_para.add_run("Giải thích: ").italic = True
                    explain_para.add_run(q.get('explanation'))
            
            # Thêm dòng trống giữa các câu
            doc.add_paragraph("")
        
        # Lưu file
        full_path = output_dir / filename
        doc.save(str(full_path))
        logger.info(f"Đã lưu file: {full_path}")
        
    except ImportError:
        logger.error("Không thể import thư viện python-docx. Vui lòng cài đặt: pip install python-docx")
        raise
    except Exception as e:
        logger.error(f"Lỗi khi xuất file docx: {str(e)}")
        raise

def resolve_book_content(ten_sach: str, noi_dung_sach: str) -> tuple[str, str]:
    """
    Tự động tìm kiếm và lấy nội dung sách từ database nếu cần thiết.
    
    Args:
        ten_sach: Tên sách cần tìm
        noi_dung_sach: Nội dung sách hoặc có thể là tên sách
        
    Returns:
        tuple(actual_content, source_info): Nội dung thực và thông tin nguồn
    """
    try:
        from tools.book_search import search_by_content
        
        # Trường hợp 1: Có tên sách rõ ràng
        if ten_sach and ten_sach.strip():
            logger.info(f"Tìm kiếm sách theo tên: {ten_sach}")
            books = search_by_content(ten_sach.strip())
            if books:
                book_info = books[0]  # Lấy kết quả đầu tiên
                content = book_info[2]  # Content column
                book_name = book_info[1]  # BookName column
                category = book_info[3]  # CategoryName column
                logger.info(f"Tìm thấy sách: {book_name} (Thể loại: {category})")
                return content, f"Sách: {book_name} (Thể loại: {category})"
            else:
                logger.warning(f"Không tìm thấy sách với tên: {ten_sach}")
        
        # Trường hợp 2: noi_dung_sach có thể là tên sách (ngắn, ít dấu chấm)
        if noi_dung_sach and noi_dung_sach.strip():
            content = noi_dung_sach.strip()
            
            # Heuristic để phát hiện có thể là tên sách:
            # - Độ dài < 100 ký tự
            # - Ít hơn 3 dấu chấm (không phải đoạn văn)
            # - Không chứa nhiều ký tự xuống dòng
            is_possibly_book_name = (
                len(content) < 100 and
                content.count('.') < 3 and
                content.count('\n') < 3 and
                not content.startswith(('Nội dung', 'Chương', 'Bài'))
            )
            
            if is_possibly_book_name:
                logger.info(f"Phát hiện có thể là tên sách: {content}")
                books = search_by_content(content)
                if books:
                    book_info = books[0]
                    book_content = book_info[2]
                    book_name = book_info[1]
                    category = book_info[3]
                    logger.info(f"Tìm thấy sách: {book_name} (Thể loại: {category})")
                    return book_content, f"Sách: {book_name} (Thể loại: {category})"
                else:
                    logger.info(f"Không tìm thấy sách, sử dụng như nội dung thông thường")
            
            # Sử dụng như nội dung thông thường
            return content, "Nội dung tùy chỉnh"
        
        # Trường hợp 3: Không có nội dung nào
        return "", "Không có nội dung tham khảo"
        
    except ImportError:
        logger.error("Không thể import search_by_content từ book_search")
        return noi_dung_sach, "Nội dung tùy chỉnh (lỗi import)"
    except Exception as e:
        logger.error(f"Lỗi khi tìm kiếm sách: {str(e)}")
        return noi_dung_sach, "Nội dung tùy chỉnh (lỗi tìm kiếm)"

@tool
def question_generator_tool(
    loai_bode: str,
    so_cau: int,
    chu_de: str = "",
    noi_dung_sach: str = "",
    ten_sach: str = ""
) -> str:
    """
    Tạo bộ đề kiểm tra từ yêu cầu của người dùng.
    Tự động tìm kiếm sách từ database nếu được cung cấp tên sách.
    
    Args:
        loai_bode: Loại bộ đề cần tạo ('trắc nghiệm' hoặc 'tự luận')
        so_cau: Số câu hỏi cần tạo (số nguyên dương)
        chu_de: Chủ đề của bộ đề (tùy chọn)
        noi_dung_sach: Nội dung sách hoặc tài liệu tham khảo (tùy chọn)
        ten_sach: Tên sách cần tìm trong database (tùy chọn)
    
    Returns:
        Chuỗi JSON chứa kết quả tạo bộ đề
    """
    import json
    
    try:
        # Tự động tìm kiếm và lấy nội dung sách nếu cần
        actual_content, source_info = resolve_book_content(ten_sach, noi_dung_sach)
        
        # Chuẩn bị request
        request = {
            "loai_bode": loai_bode,
            "so_cau": so_cau,
            "chu_de": chu_de,
            "noi_dung_sach": actual_content,
            "source_info": source_info  # Thêm thông tin nguồn
        }
        
        # Gọi hàm tạo bộ đề
        result = create_question_set(request)
        
        # Tạo response thân thiện cho người dùng
        if result['status'] == 'success':
            response = {
                "status": "thành công",
                "message": f"Đã tạo thành công bộ đề {loai_bode} với {result['question_count']} câu hỏi.",
                "source": source_info,
                "file_path": f"output/{result['docx_file']}",
                "questions_preview": [
                    {
                        "question": q.get('question', ''),
                        "type": q.get('question_type', ''),
                        "has_choices": bool(q.get('choices'))
                    } for q in result['questions'][:3]  # Chỉ hiện 3 câu đầu
                ]
            }
            
        elif result['status'] == 'need_more_info':
            response = {
                "status": "cần thêm thông tin",
                "message": "Cần cung cấp thêm thông tin:",
                "missing_info": result['questions']
            }
            
        elif result['status'] == 'validation_error':
            response = {
                "status": "lỗi validation",
                "message": "Dữ liệu đầu vào không hợp lệ:",
                "errors": result['errors']
            }
            
        else:
            response = {
                "status": "lỗi",
                "message": result.get('message', 'Đã xảy ra lỗi không xác định')
            }
        
        return json.dumps(response, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"Lỗi trong question_generator_tool: {str(e)}")
        error_response = {
            "status": "lỗi",
            "message": f"Đã xảy ra lỗi: {str(e)}"
        }
        return json.dumps(error_response, ensure_ascii=False, indent=2)
